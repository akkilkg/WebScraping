from bs4 import BeautifulSoup
import requests
from langdetect import detect
from urllib.parse import urlparse
from docx import Document
import os 


#Scrapaing Title and Content
def scrape_title_and_content(url):
    content = ""
    
    try:
        source = requests.get(url)
        source.raise_for_status()

        soup = BeautifulSoup(source.text,'html.parser')
        article = soup.find('div', class_ = "mw-content-container")
        
        
        if article:
            title = article.find('span',class_ = "mw-page-title-main").text
            contents = article.find_all('p')
            para = " ".join([content.get_text(strip = True) for content in contents])
            language = detect(para)
            indic_language_codes = [
                'hi', 'ta', 'te', 'bn', 'kn', 'gu', 'pa', 'ml', 'mr', 
                'ur', 'or', 'as', 'sa', 'kok', 'mai', 'ne', 'sd'
            ]
            
            if language in indic_language_codes:
                content += f"Title: {title}"
                content += f"Content: {para}"
            
            else:
                print("Not an Indic Language.")
        
        else:
            print("No content available.")
            
    except Exception as e:
        print(e)
    
    
    return content


#Scraping See Also Links From the Given Url
def scrape_page(url):
    
    parsed_url = urlparse(url)
    lang_code = parsed_url.netloc.split('.')[0]
    
    r = requests.get(url)
    soup = BeautifulSoup(r.text, 'html.parser')
    see_also = soup.find(id=["மேலும்_பார்க்க","इन्हें_भी_देखें","ఇవి_కూడా_చూడండి",'இவற்றையும்_பார்க்கவும்','इहो_देखल_जाय','লগতে_চাওক','আরও_দেখুন','આ_પણ_જુઓ','हे_सुद्धा_पहा','ਇਹ_ਵੀ_ਦੇਖੋ','इमानि_अपि_पश्यन्तु','مزید_دیکھیے'])
    
    links_to_follow = []
    
    if see_also:
        for element in see_also.find_all_next():
                # Stop if we reach the next section (h2)
            if element.name == 'h2':
                    break
                # Check for <ul> tags only 
            if element.name == 'ul':
                for link in element.find_all('a'): 
                    href = link.get('href')
                    if href:
                        full_url = f'https://{lang_code}.wikipedia.org{href}'
                        print(f'See Also: {full_url}')
                        links_to_follow.append(full_url)        
    
    else:
        print("No See Also Section Available")
    
    
    return links_to_follow


#Scrape Content From the Links Of See Also Section 
def scrape_see_also_section(url):
    content = " "
    
    see_also_links = scrape_page(url)
    
    for link in see_also_links:
        content += scrape_title_and_content(link)
    
    return content
    

#Saving the Scraped Content From the Url to a Word Document
def doc_save(url):
    
    doc = Document()
    doc.add_heading('Website Content', level=1)
    
    
    main_content = scrape_title_and_content(url)
    doc.add_paragraph(main_content)

    
    see_also_content = scrape_see_also_section(url)
    
    if see_also_content:
        doc.add_heading("See Also Content: ",level=2)
        doc.add_paragraph(see_also_content)

    word_doc = "Testing.docx"
    doc.save(word_doc)
    
    return word_doc

#Opening Word Document
def open_doc(path):
    if os.name == 'nt':
        os.startfile(path)


if __name__ == "__main__":
    
    url =input("Enter Wikipedia URL: ")
    word_doc_path = doc_save(url)

    if word_doc_path:
        print(f"Document Created: {word_doc_path}")
        open_doc(word_doc_path)
    else:
        print("Failed")

#url = "https://ta.wikipedia.org/wiki/%E0%AE%AF%E0%AE%AA%E0%AF%8D%E0%AE%AA%E0%AE%BE%E0%AE%A9%E0%AF%8D"
#scrape_title_and_content(url)
#scrape_page(url)
#scrape_see_also_section(url)
#doc_save(word_doc,url)